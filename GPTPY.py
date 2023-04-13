import os
import openai
import asyncio
from docx import Document
from openpyxl import load_workbook
import pandas as pd
from git import Repo
from base256 import encode as b256_encode, decode as b256_decode
from concurrent.futures import ThreadPoolExecutor

class GPTContentReviewer:
    def __init__(s, model_review, model_code):
        s.model_review = model_review
        s.model_code = model_code
        s.conversation_history = []

    async def gpt_query_with_context(s, question, model):
        role = 'role'
        content = 'content'
        s.conversation_history.append({"role": "user", "content": question})
        encoded_conversation_history = [b256_encode(str(msg)) for msg in s.conversation_history]
        response = await openai.ChatCompletion.create(
            model=model,
            messages=encoded_conversation_history,
            max_tokens=1000,
            n=1,
            stop=None,
            temperature=0.5,
        )
        answer = b256_decode(response.choices[0].message[content]).strip()
        s.conversation_history.append({role: 'system', content: answer})
        return answer

    def process_files_in_folder(s, folder_path):
        with ThreadPoolExecutor() as executor:
            for root, _, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    executor.submit(s.review_file, file_path)

    def review_file(s, file_path):
        asyncio.run(s._review_file(file_path))

    async def _review_file(s, file_path):
        with open(file_path, 'r') as f:
            content = f.read()
        review_question = f"Review the following content from the file '{file_path}': {content}"
        review_result = await s.gpt_query_with_context(review_question, s.model_review)
        task_description = "your coding task description here"
        generated_code = await s.process_coding_task(task_description)
        return generated_code

    async def process_coding_task(s, task_description):
        question = f"Write code for the following task: {task_description}"
        result = await s.gpt_query_with_context(question, s.model_code)
        return result

    async def process_local_repo(s, repo_path, output_path):
        repo = Repo(repo_path)
        files = [item.a_path for item in repo.index.diff(None)]
        for file in files:
            if file.endswith(('.docx', '.xlsx', '.csv')):
                file_type = os.path.splitext(file)[1][1:]
                file_path = os.path.join(repo.working_tree_dir, file)
                with open(file_path, 'r') as f:
                    content = f.read()
                task_description = "your coding task description here"
                generated_code = await s.review_and_generate_code(content, task_description)
                output_file_name = os.path.splitext(file)[0] + '_generated_code.py'
                output_file_path = os.path.join(output_path, output_file_name)
                with open(output_file_path, 'w') as output_file:
                    output_file.write(generated_code)
                repo.index.add(output_file_path)
                repo.index.commit(f"Add generated code for {file}")

    async def process_docx_xlsx_files(s, file_path):
        file_type = os.path.splitext(file_path)[1][1:]
        if file_type == "docx":
            document = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in document.paragraphs])
            rewrite_text_question = f"Rewrite the following docx content while retaining its original meaning: {text}"
            rewritten_text = await s.gpt_query_with_context(rewrite_text_question, s.model_review)
            new_document = Document()
            for paragraph in rewritten_text.split("\n"):
                new_document.add_paragraph(paragraph)
                new_file_path = os.path.splitext(file_path)[0] + "_rewritten.docx"
                new_document.save(new_file_path)
        elif file_type == "xlsx":
            workbook = load_workbook(file_path)
            sheet = workbook.active
            data = []
            for row in sheet.iter_rows(values_only=True):
                data.append(row)
            df = pd.DataFrame(data[1:], columns=data[0])
            rewrite_xlsx_question = f"Rewrite the following xlsx content while retaining its original meaning: {df.to_csv(index=False)}"
            rewritten_csv_data = await s.gpt_query_with_context(rewrite_xlsx_question, s.model_review)
            new_df = pd.read_csv(pd.StringIO(rewritten_csv_data))
            new_file_path = os.path.splitext(file_path)[0] + "_rewritten.xlsx"
            new_df.to_excel(new_file_path, index=False)
        return new_file_path

    async def review_eye_tracking_data(s, file):
        df = pd.read_csv(file)
        review_question = f"Review the following CSV file containing eye-tracking data/floats: {df.to_csv(index=False)}"
        review_result = await s.gpt_query_with_context(review_question, s.model_review)
        return review_result

    async def general_chat(s, message):
        chat_response = await s.gpt_query_with_context(message, s.model_review)
        return chat_response


openai.api_key = "your_openai_api_key_here"
model_review = "gpt-3.5-turbo"
model_code = "code-davinci-002"
reviewer = GPTContentReviewer(model_review, model_code)

async def main():
    input_directory = "//tedco/Documents/word_input"
    output_directory = "//tedco/Documents/html_output"
    local_repo_directory = "/mnt/c/Projects/HoloInventory/holoInventory"
    csv_input_directory = "//tedco/Documents/csv_output"
    csv_output_directory = "//tedco/Documents/csv_output"
    chat_message = "Tell me a joke."
    chat_response = await reviewer.general_chat(chat_message)
    print(f"Chat response: {chat_response}")
    await reviewer.process_local_repo(local_repo_directory, output_directory)
    print("Local Git repository files processed successfully.")


asyncio.run(main())

