import os
import openai
import asyncio
from docx import Document
from openpyxl import load_workbook
import pandas as pd
from git import Repo
from base256 import encode as b256_encode, decode as b256_decode
from concurrent.futures import ThreadPoolExecutor
from gptgui import GPTContentReviewer

class GPTContentReviewerApp:
    def __init__(self):
        self.reviewer = GPTContentReviewer("gpt-3.5-turbo", "code-davinci-002", "text-davinci-002", "text-davinci-002")
        s.model_chat = model_chat
        s.model_code = model_code
        s.model_research = model_research
        s.model_recommend = model_recommend
        s.conversation_history = []

    async def gpt_query_with_context(s, question, model):
        role = 'role'
        content = 'content'
        s.conversation_history.append({'role': 'user', 'content': question})
        encoded_conversation_history = [b256_encode(str(msg)) for msg in s.conversation_history]
        response = await openai.ChatCompletion.create(
            model=model,
            messages=encoded_conversation_history,
            max_tokens=1000,
            n=1,
            stop=None,
            temperature=0.7
        )
        answer = b256_decode(response.choices[0].message[content]).strip()
        s.conversation_history.append({role: 'system', content: answer})
        return answer

    def process_files_in_folder(s, folder_path):
        with ThreadPoolExecutor() as executor:
            for root, _, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    executor.submit(s.review_file, file_path)


    def review_file(s, file_path):
        asyncio.run(s._review_file(file_path))

    async def _review_file(s, file_path):
        with open(file_path, 'r') as f:
            content = f.read()
        review_question = f"Review the following content from the file '{file_path}': {content}"
        review_result = await s.gpt_query_with_context(review_question, s.model_review)
        task_description = "your coding task description here"
        generated_code = await s.process_coding_task(task_description)
        return generated_code

    async def process_task(s, task_description):
        recommended_model = await s.recommend_best_approach(task_description)
        research_results = await s.research_related_topic(task_description)
        generated_code = await s.process_coding_task(task_description, recommended_model)

        # Combine all results in human-readable format
        final_output = f"Recommended Model: {recommended_model}\n\nResearch Results:\n{research_results}\n\nGenerated Code:\n{generated_code}"
        human_readable_output = await s.gpt_query_with_context(final_output, s.model_chat)

        return human_readable_output

    async def process_coding_task(s, task_description, model):
        question = f"Write code for the following task: {task_description}"
        result_encoded = await s.gpt_query_with_context(question, model)
        return result_encoded

    async def process_local_repo(s, repo_path, output_path):
        repo = Repo(repo_path)
        files = [item.a_path for item in repo.index.diff(None)]
        for file in files:
            if file.endswith(('.docx', '.xlsx', '.csv')):
                file_type = os.path.splitext(file)[1][1:]
                file_path = os.path.join(repo.working_tree_dir, file)
                with open(file_path, 'r') as f:
                    content = f.read()
                task_description = "your coding task description here"
                generated_code = await s.review_and_generate_code(content, task_description)
                output_file_name = os.path.splitext(file)[0] + '_generated_code.py'
                output_file_path = os.path.join(output_path, output_file_name)
                with open(output_file_path, 'w') as output_file:
                    output_file.write(generated_code)
                repo.index.add(output_file_path)
                repo.index.commit(f"Add generated code for {file}")

    async def process_docx_xlsx_files(s, file_path):
        file_type = os.path.splitext(file_path)[1][1:]
        if file_type == "docx":
            document = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in document.paragraphs])
            rewrite_text_question = f"Rewrite the following docx content while retaining its original meaning: {text}"
            rewritten_text = await s.gpt_query_with_context(rewrite_text_question, s.model_review)
            new_document = Document()
            for paragraph in rewritten_text.split("\n"):
                new_document.add_paragraph(paragraph)
            new_file_path = os.path.splitext(file_path)[0] + "_rewritten.docx"
            new_document.save(new_file_path)
        elif file_type == "xlsx":
            workbook = load_workbook(file_path)
            sheet = workbook.active
            data = []
            for row in sheet.iter_rows(values_only=True):
                data.append(row)
            df = pd.DataFrame(data[1:], columns=data[0])
            rewrite_xlsx_question = f"Rewrite the following xlsx content while retaining its original meaning: {df.to_csv(index=False)}"
            rewritten_csv_data = await s.gpt_query_with_context(rewrite_xlsx_question, s.model_review)
            new_df = pd.read_csv(pd.StringIO(rewritten_csv_data))
            new_file_path = os.path.splitext(file_path)[0] + "_rewritten.xlsx"
            new_df.to_excel(new_file_path, index=False)
        return new_file_path
        
    async def research_related_topic(s, topic):
        question = f"Research the topic: {topic}"
        result_encoded = await s.gpt_query_with_context(question, s.model_research)
        return result_encoded

    async def recommend_best_approach(s, problem_description):
        question = f"Recommend the best approach for the following problem: {problem_description}"
        result_encoded = await s.gpt_query_with_context(question, s.model_recommend)
        return result_encoded
    async def review_eye_tracking_data(s, file):
        df = pd.read_csv(file)
        review_question = f"Review the following CSV file containing eye-tracking data/floats: {df.to_csv(index=False)}"
        review_result = await s.gpt_query_with_context(review_question, s.model_review)
        return review_result

    async def general_chat(self, message):
        chat_response = await self.reviewer.gpt_query_with_context(message, self.reviewer.model_chat)
        return chat_response
        
    async def chat_based_planning(s, user_message):
            planner_role = 'planner'
            gpt_role = 'gpt'
            user_role = 'user'
            s.conversation_history.append({planner_role: [], gpt_role: [], user_role: []})
            while True:
                try:
                    # User sends a message to GPT
                    s.conversation_history[-1][user_role].append(user_message)
                    # GPT relays the AI understanding to the planner
                    ai_understanding = await s.gpt_query_with_context(user_message, s.model_review)
                    s.conversation_history[-1][gpt_role].append(ai_understanding)
                    # Planner responds with an updated plan to GPT and Python
                    updated_plan = await s.gpt_query_with_context(ai_understanding, s.model_code)
                    s.conversation_history[-1][planner_role].append(updated_plan)
                    # GPT reads the response from the planner and sends a message to the user
                    gpt_response = await s.gpt_query_with_context(updated_plan, s.model_review)
                    s.conversation_history[-1][gpt_role].append(gpt_response)
                    # Check if the user wants to continue planning or end the conversation
                    user_message = input(f"{gpt_response}\n\nContinue planning? (yes/no):")
                    if user_message.lower() == 'no':
                        break
                except Exception as e:
                    print(f"An error occurred: {e}")
                    break
            # Encode the final plan summary and return it
            final_plan_summary = b256_encode("\n".join(s.conversation_history[-1][planner_role]))
            return final_plan_summary

    async def ai_wrangling(s, user_message):
        wrangled_message = await s.gpt_query_with_context(user_message, s.model_review)
        return wrangled_message

    async def ai_rerouting(s, user_message, target_model):
        rerouted_message = await s.gpt_query_with_context(user_message, target_model)
        return rerouted_message

# Additional roles
class Decompressor:
    def decompress(self, encoded_data):
        decoded_data = base256.decode(encoded_data)
        return decoded_data

class Indexer:
    def index_tokens(self, data):
        token_count = len(openai.api.encode(data))
        return token_count

# Check UI input for base256 encoding and token usage
def check_ui_input(ui_input):
    decompressor = Decompressor()
    indexer = Indexer()

    encoded_data = base256.encode(ui_input)
    token_count_encoded = indexer.index_tokens(encoded_data)

    if token_count_encoded > MAX_TOKENS:
        return False, "Encoded input exceeds token limit."

    decoded_data = decompressor.decompress(encoded_data)
    token_count_decoded = indexer.index_tokens(decoded_data)

    if token_count_decoded > MAX_TOKENS:
        return False, "Decoded input exceeds token limit."

    return True, "Input is within token limit."

openai.api_key = 'your_openai_api_key_here'
model_review = 'gpt-3.5-turbo'
model_code = 'code-davinci-002'
reviewer = GPTContentReviewer(model_review, model_code)


async def main():
    # Initialize GPTContentReviewer with your specific model names
    content_reviewer = GPTContentReviewer("gpt-3.5-turbo", "code-davinci-002", "text-davinci-002", "text-davinci-002")

    # Chat-based planning
    user_message = "your_initial_planning_message_here"
    final_plan_summary = await reviewer.chat_based_planning(user_message)
    print(f"Final plan summary (base)")


    # AI wrangling
    user_message = "your_message_to_be_wrangled_here"
    wrangled_message = await reviewer.ai_wrangling(user_message)
    print(f"Wrangled message: {wrangled_message}")

    # AI rerouting
    user_message = "your_message_to_be_rerouted_here"
    target_model = "code-davinci-002"  # or any other model
    rerouted_message = await reviewer.ai_rerouting(user_message, target_model)
    print(f"Rerouted message: {rerouted_message}")

    input_directory = "//tedco/Documents/word_input"
    output_directory = "//tedco/Documents/html_output"
    local_repo_directory = "/mnt/c/Projects/HoloInventory/holoInventory"
    csv_input_directory = "//tedco/Documents/csv_output"
    csv_output_directory = "//tedco/Documents/csv_output"
    chat_message = "Tell me a joke."
    chat_response = await reviewer.general_chat(chat_message)
    print(f"Chat response: {chat_response}")
    
    task_description = "Create a script to optimize and synergize everything."
    result = await content_reviewer.process_task(task_description)

    print(result)

    await reviewer.process_local_repo(local_repo_directory, output_directory)
    print("Local Git repository files processed successfully.")

    asyncio.run(main())

