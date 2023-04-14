import os
import openai
from docx import Document
from openpyxl import load_workbook
import pandas as pd
from git import Repo
from base256 import encode as b256_encode, decode as b256_decode


class GPTContentReviewer:
    """
    This class serves as the main interface for interacting with GPT-3 models and provides
    various methods to process files, documents, code, and perform research.
    """
    def __init__(self, model_chat, model_code, model_research, model_recommend):
        self.model_chat = model_chat
        self.model_code = model_code
        self.model_research = model_research
        self.model_recommend = model_recommend
        self.model_recommend = model_recommend
        self.conversation_history = []

    def list_files_in_repo(self, local_repo_path):
        repo = Repo(local_repo_path)
        default_branch = repo.head.reference  # Use the HEAD reference to get the default branch
        file_list = [item.path for item in default_branch.commit.tree.traverse() if item.type == "blob"]
        return file_list

    def gpt_query_with_context(self, query, model, max_tokens=1000, role='user', history='conversation', is_chat_based=True):
        if self.check_token_amount(query) > max_tokens:
            return self.process_large_input(query, model, max_tokens, role, history, is_chat_based)

        if is_chat_based:
            hist = self.conversation_history
            hist.append({'role': role, 'content': query})
            response = openai.ChatCompletion.create(model=model, messages=hist[-10:], max_tokens=max_tokens, n=1, stop=None, temperature=0.7)
            answer_text = response.choices[0].message['content'].strip()
            hist.append({'role': 'system', 'content': answer_text})
        else:
            prompt = f"{query}:"
            response = openai.Completion.create(engine=model, prompt=prompt, max_tokens=max_tokens, n=1, stop=None, temperature=0.7)
            answer_text = response.choices[0].text.strip()

        return answer_text

    def process_large_input(self, input_text, model, max_tokens=1000, role='user', history='conversation', is_chat_based=True):
        chunk_size = max_tokens - 10  # Leaving some tokens for additional context or instructions

        input_chunks = [input_text[i:i + chunk_size] for i in range(0, len(input_text), chunk_size)]

        output_chunks = []

        for chunk in input_chunks:
            output_chunk = self.gpt_query_with_context(chunk, model, max_tokens=max_tokens, role=role, history=history, is_chat_based=is_chat_based)
            output_chunks.append(output_chunk)
        return "".join(output_chunks)

    def general_chat(self, message):
        chat_response = self.gpt_query_with_context(message, self.model_chat, is_chat_based=True)
        return chat_response

    def check_token_amount(self, encoded_message):
        token_amount = len(openai.api.encoders.default.encode(encoded_message))
        return token_amount

    def process_task(self, task_description):
        recommended_model = self.recommend_best_approach(task_description)
        research_results = self.research_related_topic(task_description)
        generated_code = self.process_coding_task(task_description, recommended_model)
        final_output = f"RecommendedModel:{recommended_model}\nResearchResults:{research_results}\nGeneratedCode:{generated_code}"
        human_readable_output = self.general_chat(final_output)
        return human_readable_output

    def process_docx_xlsx_files(self, file_path):
        file_type = os.path.splitext(file_path)[1][1:]
        if file_type == "docx":
            document = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in document.paragraphs])
            rewrite_text_question = f"Rewrite the following docx content while retaining its original meaning: {text}"
    
            rewritten_text = self.gpt_query_with_context(rewrite_text_question, self.model_chat, is_chat_based=True)
            new_document = Document()
            for paragraph in rewritten_text.split("\n"):
                new_document.add_paragraph(paragraph)
            new_file_path = os.path.splitext(file_path)[0] + "_rewritten.docx"
            new_document.save(new_file_path)
        elif file_type == "xlsx":
            workbook = load_workbook(file_path)
            sheet = workbook.active
            data = []
            for row in sheet.iter_rows(values_only=True):
                data.append(row)
            df = pd.DataFrame(data[1:], columns=data[0])
            rewrite_xlsx_question = f"Rewrite the following xlsx content while retaining its original meaning: {df.to_csv(index=False)}"
            rewritten_csv_data = self.gpt_query_with_context(rewrite_xlsx_question, self.model_chat, is_chat_based=True)
            new_df = pd.read_csv(pd.StringIO(rewritten_csv_data))
            new_file_path = os.path.splitext(file_path)[0] + "_rewritten.xlsx"
            new_df.to_excel(new_file_path, index=False)
        return new_file_path

    def process_coding_task(self, task_description, model):
        question = f"Write code for the following task: {task_description}"
        result_encoded = self.gpt_query_with_context(question, model, is_chat_based=True)
        return result_encoded

    def research_related_topic(self, topic):
        question = f"Research the topic: {topic}"
        result_encoded = self.gpt_query_with_context(question, self.model_research, is_chat_based=False)
        return result_encoded

    def recommend_best_approach(self, problem_description):
        research_results = self.research_related_topic(problem_description)
        question = f"Recommend the best approach for the following problem, considering the research: {problem_description}\n\nResearchResults:\n{research_results}"
        result_encoded = self.gpt_query_with_context(question, self.model_research, is_chat_based=False)
        return result_encoded

    def review_eye_tracking_data(self, file):
        df = pd.read_csv(file)
        review_question = f"Review the following CSV file containing eye-tracking data/floats: {df.to_csv(index=False)}"
        review_result = self.gpt_query_with_context(review_question, self.model_chat, is_chat_based=True)
        return review_result

    def chat_based_planning(self, user_message):
        planner_role = 'planner'
        gpt_role = 'gpt'
        user_role = 'user'
        self.conversation_history.append({planner_role: [], gpt_role: [], user_role: []})
        while True:
            try:
                self.conversation_history[-1][user_role].append(user_message)
                ai_understanding = self.gpt_query_with_context(user_message, self.model_chat, is_chat_based=True)
                self.conversation_history[-1][gpt_role].append(ai_understanding)
                updated_plan = self.gpt_query_with_context(ai_understanding, self.model_code, is_chat_based=True)
                self.conversation_history[-1][planner_role].append(updated_plan)
                gpt_response = self.gpt_query_with_context(updated_plan, self.model_chat, is_chat_based=True)
                self.conversation_history[-1][gpt_role].append(gpt_response)
                user_message = input(f"{gpt_response}\n\nContinue planning? (yes/no):")
                if user_message.lower() == 'no':
                    break
            except Exception as e:
                print(f"An error occurred: {e}")
                break
        final_plan_summary = b256_encode('\n'.join(self.conversation_history[-1][planner_role]))
        return final_plan_summary

    def ai_wrangling(self, user_message):
        wrangled_message = self.gpt_query_with_context(user_message, self.model_chat, is_chat_based=True)
        return wrangled_message

    def ai_rerouting(self, user_message, target_model):
        rerouted_message = self.gpt_query_with_context(user_message, target_model, is_chat_based=True)
        return rerouted_message

class Recommender(GPTContentReviewer):
    def recommend_based_on_eye_tracking_data(self, file):
        review_result = self.review_eye_tracking_data(file)
        recommendation_question = f"Provide recommendations based on the following eye-tracking data review:\n\n{review_result}"
        recommendation_result = self.gpt_query_with_context(recommendation_question, self.model_recommend, is_chat_based=False)
        return recommendation_result

    def recommend_next_step_in_planning(self, user_message):
        chat_result = self.chat_based_planning(user_message)
        decoded_chat_result = b256_decode(chat_result)
        recommendation_question = f"Recommend the next step in planning based on the following chat-based planning summary:\n\n{decoded_chat_result}"
        recommendation_result = self.gpt_query_with_context(recommendation_question, self.model_recommend, is_chat_based=False)
        return recommendation_result


openai.api_key = "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"
openai.api_base = "https://api.openai.com/v1"  # Add this line
content_reviewer_app = GPTContentReviewer("gpt-3.5-turbo", "code-davinci-002", "text-davinci-002", "text-davinci-002")

def main():
    while True:
        print('\nPlease choose an option:')
        print('A. Process Word documents or Excel files using GPT (docx/xlsx review and rewrite)')
        print('B. Process local repo coding review and rewrite')
        print('C. Process CSV file for eye tracking data assessment')
        print('D. Direct interaction with GPT (without any coding or review context, but with context history still in memory if applicable)')
        print('E. Chat-based planning with file output creation')
        print('Q. Quit')

        choice = input('Enter your choice:').lower()

        if choice == 'a':
            file_path = input('Enter the path to the file (docx or xlsx):')
            new_file_path = content_reviewer_app.process_docx_xlsx_files(file_path)
            print(f"Processed file saved at: {new_file_path}")
        elif choice == 'b':
            local_repo_path = input("Enter the path to the local repository:")
            file_list = content_reviewer_app.list_files_in_repo(local_repo_path)
            file_list_str = "\n".join(file_list)
            task_description = f"What files do you have in the repo?\n\n{file_list_str}"
            output = content_reviewer_app.process_task(task_description)
            print(f"Output: {output}")

        elif choice == 'c':
            file_path = input('Enter the path to the CSV file containing eye-tracking data:')
            result = content_reviewer_app.recommend_based_on_eye_tracking_data(file_path)
            print(f"Recommendation: {result}")
        elif choice == 'd':
            user_message = input('Enter your message for GPT interaction:')
            chat_response = content_reviewer_app.general_chat(user_message)
            print(f"Chat response: {chat_response}")
        elif choice == 'e':
            user_message = input("Enter your initial message for chat-based planning:")
            planning_summary = content_reviewer_app.chat_based_planning(user_message)
            decoded_summary = b256_decode(planning_summary)
            print(f"Planning summary: {decoded_summary}")
            file_name = input("Enter the file name for the output file:")
            with open(file_name, "w") as f:
                f.write(decoded_summary)
            print(f"File saved as: {file_name}")
        elif choice == 'q':
            break
        else:
            print('Invalid choice. Please try again.')

if __name__ == '__main__':
    main()

