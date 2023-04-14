import os
import openai
from docx import Document
from openpyxl import load_workbook
import pandas as pd
from git import Repo
from base256 import encode as b256_encode, decode as b256_decode


class GPTContentReviewer:
    def __init__(self, model_chat, model_code, model_research, model_recommend):
        self.model_chat = model_chat
        self.model_code = model_code
        self.model_research = model_research
        self.model_recommend = model_recommend
        self.conversation_history = []

    def gpt_query_with_context(self, query, model, max_tokens=1000, role='user', history='conversation'):
        hist = self.conversation_history
        hist.append({'role': role, 'content': query})
        response = openai.ChatCompletion.create(
            model=model,
            messages=hist[-10:],
            max_tokens=max_tokens,
            n=1,
            stop=None,
            temperature=0.7
        )
        answer_text = response.choices[0].message['content'].strip()
        hist.append({'role': 'system', 'content': answer_text})
        return answer_text

    def general_chat(self, message):
        #encoded_message = b256_encode(message)
        chat_response = self.gpt_query_with_context(message, self.model_chat)
        #decoded_response = b256_decode(chat_response)
        return chat_response

    def check_token_amount(self, encoded_message):
        token_amount = len(openai.api.encoders.default.encode(encoded_message))
        return token_amount

    def process_task(self, task_description):
        recommended_model = self.recommend_best_approach(task_description)
        research_results = self.research_related_topic(task_description)
        generated_code = self.process_coding_task(task_description, recommended_model)
        final_output = f"RecommendedModel: {recommended_model}\n\nResearchResults:\n{research_results}\n\nGeneratedCode:\n{generated_code}"
        human_readable_output = self.general_chat(final_output)
        return human_readable_output

    def process_docx_xlsx_files(self, file_path):
        file_type = os.path.splitext(file_path)[1][1:]
        if file_type == "docx":
            document = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in document.paragraphs])
            rewrite_text_question = f"Rewrite the following docx content while retaining its original meaning: {text}"
    
            rewritten_text = self.gpt_query_with_context(rewrite_text_question, self.model_chat)
            new_document = Document()
            for paragraph in rewritten_text.split("\n"):
                new_document.add_paragraph(paragraph)
            new_file_path = os.path.splitext(file_path)[0] + "_rewritten.docx"
            new_document.save(new_file_path)
        elif file_type == "xlsx":
            workbook = load_workbook(file_path)
            sheet = workbook.active
            data = []
            for row in sheet.iter_rows(values_only=True):
                data.append(row)
            df = pd.DataFrame(data[1:], columns=data[0])
            rewrite_xlsx_question = f"Rewrite the following xlsx content while retaining its original meaning: {df.to_csv(index=False)}"
            rewritten_csv_data = self.gpt_query_with_context(rewrite_xlsx_question, self.model_chat)
            new_df = pd.read_csv(pd.StringIO(rewritten_csv_data))
            new_file_path = os.path.splitext(file_path)[0] + "_rewritten.xlsx"
            new_df.to_excel(new_file_path, index=False)
        return new_file_path

    def process_coding_task(self, task_description, model):
        question = f"Write code for the following task: {task_description}"
        result_encoded = self.gpt_query_with_context(question, model)
        return result_encoded

    def research_related_topic(self, topic):
        question = f"Research the topic: {topic}"
        result_encoded = self.gpt_query_with_context(question, self.model_research)
        return result_encoded

    def recommend_best_approach(self, problem_description):
        research_results = self.research_related_topic(problem_description)
        question = f"Recommend the best approach for the following problem, considering the research: {problem_description}\n\nResearchResults:\n{research_results}"
        result_encoded = self.gpt_query_with_context(question, self.model_research)
        return result_encoded

    def review_eye_tracking_data(self, file):
        df = pd.read_csv(file)
        review_question = f"Review the following CSV file containing eye-tracking data/floats: {df.to_csv(index=False)}"
        review_result = self.gpt_query_with_context(review_question, self.model_chat)
        return review_result

    def chat_based_planning(self, user_message):
        planner_role = "planner"
        gpt_role = "gpt"
        user_role = "user"
        self.conversation_history.append({planner_role: [], gpt_role: [], user_role: []})
        while True:
            try:
                self.conversation_history[-1][user_role].append(user_message)
                ai_understanding = self.gpt_query_with_context(user_message, self.model_chat)
                self.conversation_history[-1][gpt_role].append(ai_understanding)
                updated_plan = self.gpt_query_with_context(ai_understanding, self.model_code)
                self.conversation_history[-1][planner_role].append(updated_plan)
                gpt_response = self.gpt_query_with_context(updated_plan, self.model_chat)
                self.conversation_history[-1][gpt_role].append(gpt_response)
                user_message = input(f"{gpt_response}\n\nContinue planning? (yes/no):")
                if user_message.lower() == "no":
                    break
            except Exception as e:
                print(f"An error occurred: {e}")
                break
        final_plan_summary = b256_encode("\n".join(self.conversation_history[-1][planner_role]))
        return final_plan_summary

    def ai_wrangling(self, user_message):
        wrangled_message = self.gpt_query_with_context(user_message, self.model_chat)
        return wrangled_message

    def ai_rerouting(self, user_message, target_model):
        rerouted_message = self.gpt_query_with_context(user_message, target_model)
        return rerouted_message

class FlaskDataOverride:
    def __init__(self, app):
        self.app = app

    def override_folder(self, folder_path):
        self.app.config['FOLDER_PATH'] = folder_path

    def override_key(self, api_key):
        self.app.config['API_KEY'] = api_key

    def update_content_reviewer_app(self):
        folder_path = self.app.config.get('FOLDER_PATH')
        api_key = self.app.config.get('API_KEY')
        if folder_path:
            content_reviewer_app.folder_path = folder_path
        if api_key:
            content_reviewer_app.api_key = api_key

class Decompressor:
    def decompress(self, encoded_data):
        decoded_data = base256.decode(encoded_data)
        return decoded_data

class Indexer:
    def index_tokens(self, data):
        token_count = len(openai.api.encode(data))
        return token_count

class Recommender(GPTContentReviewer):
    async def recommend_based_on_eye_tracking_data(self, file):
        review_result = await self.review_eye_tracking_data(file)
        recommendation_question = f"Provide recommendations based on the following eye-tracking data review:\n\n{review_result}"
        recommendation_result = await self.gpt_query_with_context(recommendation_question, self.model_recommend)
        return recommendation_result

    async def recommend_next_step_in_planning(self, user_message):
        chat_result = await self.chat_based_planning(user_message)
        decoded_chat_result = b256_decode(chat_result)
        recommendation_question = f"Recommend the next step in planning based on the following chat-based planning summary:\n\n{decoded_chat_result}"
        recommendation_result = await self.gpt_query_with_context(recommendation_question, self.model_recommend)
        return recommendation_result

openai.api_key = "sk-"
content_reviewer_app = GPTContentReviewer("gpt-3.5-turbo", "code-davinci-002", "text-davinci-002", "text-davinci-002")
model_review = "gpt-3.5-turbo"
model_code = "code-davinci-002"


if __name__ == '__main__':
    openai.api_key = "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"
    content_reviewer_app = GPTContentReviewer("gpt-3.5-turbo", "code-davinci-002", "text-davinci-002", "text-davinci-002")

    user_message = 'Tell me a joke.'
    chat_response = content_reviewer_app.general_chat(user_message)
    print(f"Chat response: {chat_response}")

    task_description = 'Create a script to optimize and synergize everything.'
    result = content_reviewer_app.process_task(task_description)
    print(result)

    user_message_to_wrangle = 'This is an example message for AI wrangling.'
    wrangled_message = content_reviewer_app.ai_wrangling(user_message_to_wrangle)
    print(f"Wrangled message: {wrangled_message}")

    user_message_to_route = 'This is an example message for AI rerouting.'
    target_model_for_rerouting = 'code-davinci-002'
    rerouted_message = content_reviewer_app.ai_rerouting(user_message_to_route, target_model_for_rerouting)
    print(f"Rerouted message: {rerouted_message}")
